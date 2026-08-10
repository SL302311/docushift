#!/usr/bin/env python3
"""
DocuShift 转换引擎
===================
支持 PDF / DOCX / MD / HTML / 图片 之间的互转。
所有转换方法均在独立线程中调用，通过 progress_callback 回调更新 UI。

转换路由表:
  PDF  → DOCX | PNG | JPG          (pdf2docx / PyMuPDF)
  DOCX → PDF                        (docx2pdf / pypandoc)
  MD   → DOCX | HTML | PDF          (pypandoc)
  HTML → DOCX | PDF                 (pypandoc)
  PNG  → PDF                        (PyMuPDF)
  JPG  → PDF                        (PyMuPDF)
"""

from __future__ import annotations

import os
import sys
import shutil
import traceback
from typing import Callable, Optional

# ---------------------------------------------------------------------------
# 依赖探测 —— 按需导入，缺失时给出清晰提示
# ---------------------------------------------------------------------------

_import_errors: dict[str, str] = {}

try:
    from pdf2docx import Converter as _Pdf2DocxConverter
except Exception as e:
    _import_errors["pdf2docx"] = str(e)

try:
    import fitz  # PyMuPDF
except Exception as e:
    _import_errors["PyMuPDF"] = str(e)

try:
    import pypandoc
except Exception as e:
    _import_errors["pypandoc"] = str(e)

try:
    from docx2pdf import convert as _docx2pdf_convert
except Exception as e:
    _import_errors["docx2pdf"] = str(e)


# ---------------------------------------------------------------------------
# 格式支持表
# ---------------------------------------------------------------------------

SUPPORTED_FORMATS: dict[str, list[str]] = {
    ".pdf":  ["docx", "png", "jpg"],
    ".docx": ["pdf"],
    ".doc":  ["pdf"],
    ".md":   ["docx", "html", "pdf"],
    ".markdown": ["docx", "html", "pdf"],
    ".html": ["docx", "pdf"],
    ".htm":  ["docx", "pdf"],
    ".png":  ["pdf"],
    ".jpg":  ["pdf"],
    ".jpeg": ["pdf"],
    ".bmp":  ["pdf"],
    ".tiff": ["pdf"],
}

# 人类可读的格式描述
FORMAT_LABELS: dict[str, str] = {
    "docx": "Word 文档 (.docx)",
    "pdf":  "PDF 文档 (.pdf)",
    "html": "网页 (.html)",
    "png":  "PNG 图片 (.png)",
    "jpg":  "JPG 图片 (.jpg)",
}

# 输出文件扩展名映射
OUTPUT_EXTENSIONS: dict[str, str] = {
    "docx": ".docx",
    "pdf":  ".pdf",
    "html": ".html",
    "png":  ".png",
    "jpg":  ".jpg",
}


def get_output_formats(file_path: str) -> list[str]:
    """根据输入文件后缀返回可选的输出格式列表。"""
    ext = os.path.splitext(file_path)[1].lower()
    return SUPPORTED_FORMATS.get(ext, [])


def get_pandoc_path() -> Optional[str]:
    """
    查找 pandoc 可执行文件路径。
    优先级: 同目录 > _MEIPASS > 用户缓存 ~/.docushift > pypandoc > 系统 PATH。
    如均未找到，自动下载到 ~/.docushift/pandoc.exe。
    """
    # 1) 同目录或 _MEIPASS（打包场景）
    if getattr(sys, "frozen", False):
        base = sys._MEIPASS  # type: ignore
    else:
        base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    local = os.path.join(base, "pandoc.exe")
    if os.path.isfile(local):
        return local

    # 2) 用户缓存目录 (~/.docushift/pandoc.exe)
    cache_dir = os.path.join(os.path.expanduser("~"), ".docushift")
    cached = os.path.join(cache_dir, "pandoc.exe")
    if os.path.isfile(cached):
        return cached

    # 3) pypandoc 已下载版本
    try:
        return pypandoc.get_pandoc_path()
    except Exception:
        pass

    # 4) 系统 PATH
    import shutil as _shutil
    if _shutil.which("pandoc"):
        return _shutil.which("pandoc") or None

    # 5) 自动下载到用户缓存目录
    return _auto_download_pandoc(cache_dir)


def _auto_download_pandoc(target_dir: str) -> Optional[str]:
    """首次使用 MD/HTML 转换时自动下载 pandoc.exe。"""
    import urllib.request, zipfile, io as _io

    url = (
        "https://github.com/jgm/pandoc/releases/"
        "download/3.1.11/pandoc-3.1.11-windows-x86_64.zip"
    )
    os.makedirs(target_dir, exist_ok=True)
    out_path = os.path.join(target_dir, "pandoc.exe")

    print("[DocuShift] 首次下载 pandoc (约 200MB)，请稍候…")
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        resp = urllib.request.urlopen(req, timeout=120)
        data = resp.read()

        zf = zipfile.ZipFile(_io.BytesIO(data))
        for name in zf.namelist():
            if name.endswith("pandoc.exe"):
                with zf.open(name) as src, open(out_path, "wb") as dst:
                    dst.write(src.read())
                print(f"[DocuShift] pandoc 下载完成: {out_path}")
                return out_path
        return None
    except Exception:
        print("[DocuShift] pandoc 下载失败，MD/HTML 转换暂不可用")
        return None


# ---------------------------------------------------------------------------
# 转换引擎
# ---------------------------------------------------------------------------

class ConversionError(Exception):
    """转换过程中的业务异常。"""


class DocumentConverter:
    """文档格式转换引擎，线程安全（无共享可变状态）。"""

    def __init__(self) -> None:
        self._pandoc_path: Optional[str] = None  # 懒加载，首次需要时才查找

    def _get_pandoc(self) -> Optional[str]:
        """懒加载 pandoc 路径（仅 MD/HTML 转换时触发）。"""
        if self._pandoc_path is None:
            self._pandoc_path = get_pandoc_path()
        return self._pandoc_path

    # ---- 公共接口 ---------------------------------------------------------

    def convert(
        self,
        input_path: str,
        output_format: str,
        progress_callback: Optional[Callable[[int, str], None]] = None,
    ) -> list[str]:
        """
        执行单文件转换。

        Args:
            input_path:     输入文件绝对路径
            output_format:  目标格式 (docx / pdf / html / png / jpg)
            progress_callback: 回调函数 (percent: int, message: str)

        Returns:
            输出文件路径列表

        Raises:
            ConversionError: 转换失败
        """
        input_path = os.path.abspath(input_path)
        if not os.path.isfile(input_path):
            raise ConversionError(f"文件不存在: {input_path}")

        ext = os.path.splitext(input_path)[1].lower()
        output_format = output_format.lower()

        if ext not in SUPPORTED_FORMATS:
            raise ConversionError(f"不支持的输入格式: {ext}")

        available = SUPPORTED_FORMATS[ext]
        if output_format not in available:
            raise ConversionError(
                f"不支持 {ext} → .{output_format}，可选: {available}"
            )

        self._emit(progress_callback, 0, "正在准备转换…")

        try:
            # ---- 路由 -----------------------------------------------------
            if ext == ".pdf" and output_format == "docx":
                result = self._pdf_to_docx(input_path, progress_callback)
            elif ext == ".pdf" and output_format in ("png", "jpg"):
                result = self._pdf_to_images(input_path, output_format, progress_callback)
            elif ext in (".docx", ".doc") and output_format == "pdf":
                result = self._docx_to_pdf(input_path, progress_callback)
            elif ext in (".md", ".markdown") and output_format in ("docx", "html", "pdf"):
                result = self._pandoc_convert(input_path, output_format, progress_callback)
            elif ext in (".html", ".htm") and output_format in ("docx", "pdf"):
                result = self._pandoc_convert(input_path, output_format, progress_callback)
            elif ext in (".png", ".jpg", ".jpeg", ".bmp", ".tiff") and output_format == "pdf":
                result = self._image_to_pdf(input_path, progress_callback)
            else:
                raise ConversionError(f"未实现的转换: {ext} → {output_format}")

            self._emit(progress_callback, 100, "转换完成")
            return result

        except ConversionError:
            raise
        except Exception as e:
            raise ConversionError(f"{e}\n{traceback.format_exc()}") from e

    def _get_output_path(self, input_path: str, output_ext: str) -> str:
        """
        生成输出文件路径：源文件同目录下创建 DocuShift_Output 文件夹。
        若同名文件已存在，自动追加序号 (如 _1, _2) 避免静默覆盖。
        """
        src_dir = os.path.dirname(os.path.abspath(input_path))
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        out_dir = os.path.join(src_dir, "DocuShift_Output")
        os.makedirs(out_dir, exist_ok=True)

        # 检查是否已存在同名文件
        out_path = os.path.join(out_dir, base_name + output_ext)
        if not os.path.exists(out_path):
            return out_path

        # 自动追加序号避免覆盖
        counter = 1
        while True:
            out_path = os.path.join(out_dir, f"{base_name}_{counter}{output_ext}")
            if not os.path.exists(out_path):
                return out_path
            counter += 1

    # ---- PDF → DOCX -------------------------------------------------------

    def _pdf_to_docx(
        self,
        pdf_path: str,
        progress_cb: Optional[Callable[[int, str], None]],
    ) -> list[str]:
        """
        PDF → DOCX 转换。
        使用 PyMuPDF 提取结构化文本 + python-docx 重建干净文档，
        避免pdf2docx的文本框碎片化和大量留白问题。
        """
        if "PyMuPDF" in _import_errors:
            raise ConversionError("PyMuPDF 库未安装，无法执行 PDF→Word 转换")
        if "python-docx" in _import_errors:
            raise ConversionError("python-docx 库未安装，无法执行 PDF→Word 转换")

        docx_path = os.path.splitext(pdf_path)[0] + ".docx"
        docx_path = self._get_output_path(pdf_path, ".docx")
        self._emit(progress_cb, 10, "正在解析 PDF…")

        doc = fitz.open(pdf_path)
        total_pages = len(doc)

        # 文本清洗：移除 XML 不兼容的控制字符
        import re
        def clean_text(text: str) -> str:
            # 移除除 \t \n \r 外的所有控制字符
            text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
            # 合并连续空格
            text = re.sub(r' {3,}', '  ', text)
            return text.strip()

        # 第一遍：收集所有页面的文本块，统计字体大小分布以识别标题
        self._emit(progress_cb, 20, f"正在分析 {total_pages} 页内容…")
        all_blocks = []
        font_sizes = []
        for page in doc:
            page_dict = page.get_text("dict")
            for block in page_dict.get("blocks", []):
                if block.get("type") == 0:  # 文本块
                    for line in block.get("lines", []):
                        line_text = ""
                        line_size = 0
                        line_flags = 0
                        for span in line.get("spans", []):
                            line_text += span.get("text", "")
                            line_size = max(line_size, span.get("size", 12))
                            line_flags |= span.get("flags", 0)
                        line_text = clean_text(line_text)
                        if line_text:
                            all_blocks.append({
                                "text": line_text,
                                "size": round(line_size, 1),
                                "bold": bool(line_flags & 16),  # bit 4 = bold
                            })
                            font_sizes.append(round(line_size, 1))

        # 判断正文字体大小（出现频率最高的）
        if font_sizes:
            from collections import Counter
            size_counter = Counter(font_sizes)
            body_size = size_counter.most_common(1)[0][0]
            # 标题 = 明显大于正文的字体
            heading_threshold = body_size + 2.0
        else:
            body_size = 12.0
            heading_threshold = 14.0

        # 第二遍：用 python-docx 创建文档
        self._emit(progress_cb, 50, "正在生成 Word 文档…")
        from docx import Document
        from docx.shared import Pt, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        word_doc = Document()

        # 设置默认字体
        style = word_doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(body_size)

        img_count = 0
        seen_xrefs = set()  # 去重：同一 xref 只提取一次
        import io as _io
        from PIL import Image as _PILImage

        for page_idx, page in enumerate(doc):
            # 按 block 顺序处理，保持原文阅读顺序（文本和图片混合）
            page_dict = page.get_text("dict")
            for block in page_dict.get("blocks", []):
                # ---- 图片块：按出现位置插入，缩小尺寸 ----
                if block.get("type") == 1:
                    bbox = block.get("bbox", [0, 0, 0, 0])
                    w_pt = bbox[2] - bbox[0] if len(bbox) == 4 else 0
                    h_pt = bbox[3] - bbox[1] if len(bbox) == 4 else 0
                    # 过滤过小图片
                    if w_pt < 50 and h_pt < 50:
                        continue
                    # 获取 xref 去重
                    xref = block.get("image", 0) if isinstance(block.get("image"), int) else 0
                    if not xref:
                        # 从 page images 中按 bbox 匹配
                        for pi in page.get_images(full=True):
                            if pi[0] not in seen_xrefs:
                                xref = pi[0]
                                break
                    if not xref or xref in seen_xrefs:
                        continue
                    seen_xrefs.add(xref)
                    try:
                        img_bytes = doc.extract_image(xref)["image"]
                        img_stream = _io.BytesIO(img_bytes)
                        pil_img = _PILImage.open(img_stream)
                        ow, oh = pil_img.size
                        img_stream.seek(0)
                        # 根据原 PDF 中的尺寸设宽度，最大 4 英寸
                        target_w = min(4.0, max(1.5, w_pt / 72.0))
                        word_doc.add_picture(img_stream, width=Inches(target_w))
                        img_count += 1
                    except Exception:
                        pass
                    continue

                # ---- 文本块 ----
                if block.get("type") != 0:
                    continue

                block_lines = []
                block_max_size = 0
                block_bold = False
                for line in block.get("lines", []):
                    line_text = ""
                    line_size = 0
                    line_flags = 0
                    for span in line.get("spans", []):
                        line_text += span.get("text", "")
                        line_size = max(line_size, span.get("size", 12))
                        line_flags |= span.get("flags", 0)
                    line_text = clean_text(line_text)
                    if line_text:
                        block_lines.append(line_text)
                        block_max_size = max(block_max_size, round(line_size, 1))
                        if line_flags & 16:
                            block_bold = True

                if not block_lines:
                    continue

                block_text = " ".join(block_lines)
                block_text = re.sub(r' {2,}', ' ', block_text).strip()

                if block_max_size >= heading_threshold and len(block_text) < 100:
                    if block_max_size >= heading_threshold + 4:
                        word_doc.add_heading(block_text, level=1)
                    else:
                        word_doc.add_heading(block_text, level=2)
                else:
                    p = word_doc.add_paragraph(block_text)
                    if block_bold:
                        for run in p.runs:
                            run.bold = True

            # 页面间不插入分页符（让内容自然流动）
            pct = 50 + int(40 * (page_idx + 1) / total_pages)
            self._emit(progress_cb, pct, f"已处理 {page_idx + 1}/{total_pages} 页")

        doc.close()

        self._emit(progress_cb, 92, "正在保存 Word 文档…")
        word_doc.save(docx_path)
        self._emit(progress_cb, 95, f"完成（{img_count} 张图片）")
        return [docx_path]

    # ---- PDF → 图片 -------------------------------------------------------

    def _pdf_to_images(
        self,
        pdf_path: str,
        img_format: str,
        progress_cb: Optional[Callable[[int, str], None]],
    ) -> list[str]:
        if "PyMuPDF" in _import_errors:
            raise ConversionError("PyMuPDF 库未安装，无法执行 PDF→图片 转换")

        # 输出到 DocuShift_Output 文件夹
        src_dir = os.path.dirname(os.path.abspath(pdf_path))
        out_dir = os.path.join(src_dir, "DocuShift_Output")
        os.makedirs(out_dir, exist_ok=True)
        base_name = os.path.splitext(os.path.basename(pdf_path))[0]

        doc = fitz.open(pdf_path)
        total = len(doc)
        results: list[str] = []

        for i, page in enumerate(doc):
            # 2x 缩放保证清晰度
            matrix = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=matrix, alpha=False)

            out_name = f"{base_name}_p{i + 1}.{img_format}"
            out_path = os.path.join(out_dir, out_name)

            if img_format == "jpg":
                pix.pil_save(out_path, format="JPEG", quality=95)
            else:
                pix.save(out_path)

            results.append(out_path)
            pct = int(90 * (i + 1) / total)
            self._emit(progress_cb, pct, f"已导出 {i + 1}/{total} 页")

        doc.close()
        return results

    # ---- DOCX → PDF -------------------------------------------------------

    def _docx_to_pdf(
        self,
        docx_path: str,
        progress_cb: Optional[Callable[[int, str], None]],
    ) -> list[str]:
        pdf_path = self._get_output_path(docx_path, ".pdf")

        # 方案一: 直接调用 MS Word COM（最可靠，不依赖 docx2pdf 库）
        if sys.platform == "win32":
            try:
                self._emit(progress_cb, 20, "正在通过 MS Word 转换…")
                self._word_to_pdf_com(docx_path, pdf_path)
                if os.path.isfile(pdf_path):
                    return [pdf_path]
            except Exception as e:
                err = str(e)[:100]
                self._emit(progress_cb, 30, f"Word 转换失败，尝试备用方案…")

        # 方案二: docx2pdf 库（回退）
        if "docx2pdf" not in _import_errors:
            try:
                self._emit(progress_cb, 35, "正在通过 docx2pdf 转换…")
                _docx2pdf_convert(docx_path, pdf_path)
                if os.path.isfile(pdf_path):
                    return [pdf_path]
            except Exception:
                pass

        # 方案三: pypandoc（需要 pandoc + LaTeX）
        self._emit(progress_cb, 40, "正在通过 pandoc 转换…")
        result = self._pandoc_convert(docx_path, "pdf", progress_cb)
        return result

    @staticmethod
    def _word_to_pdf_com(docx_path: str, pdf_path: str) -> None:
        """
        直接通过 Win32 COM 调用 MS Word 执行 DOCX→PDF 转换。
        比 docx2pdf 库更可靠，无 tqdm/控制台依赖。
        """
        import pythoncom
        import win32com.client

        # 在当前线程初始化 COM（--windowed 打包模式必需）
        pythoncom.CoInitialize()
        word_app = None
        doc = None
        try:
            word_app = win32com.client.DispatchEx("Word.Application")
            word_app.Visible = False
            word_app.DisplayAlerts = False

            # 打开文档（只读）
            doc = word_app.Documents.Open(
                os.path.abspath(docx_path),
                ReadOnly=True,
                AddToRecentFiles=False,
            )

            # 导出为 PDF (17 = wdExportFormatPDF)
            doc.ExportAsFixedFormat(
                OutputFileName=os.path.abspath(pdf_path),
                ExportFormat=17,
                OpenAfterExport=False,
                OptimizeFor=0,   # wdExportOptimizeForPrint
                Range=0,         # wdExportAllDocument
                Item=0,          # wdExportDocumentContent
            )
        finally:
            if doc:
                doc.Close(SaveChanges=False)
            if word_app:
                word_app.Quit()
            pythoncom.CoUninitialize()

    # ---- MD / HTML → DOCX / HTML / PDF (pypandoc) -------------------------

    def _pandoc_convert(
        self,
        input_path: str,
        output_format: str,
        progress_cb: Optional[Callable[[int, str], None]],
    ) -> list[str]:
        if "pypandoc" in _import_errors:
            raise ConversionError("pypandoc 库未安装，无法执行此转换")

        out_ext = OUTPUT_EXTENSIONS[output_format]
        output_path = self._get_output_path(input_path, out_ext)

        self._emit(progress_cb, 30, f"正在通过 pandoc 转换为 {output_format.upper()}…")

        extra_args: list[str] = []
        if output_format == "pdf":
            # PDF 需要指定 PDF 引擎；优先用 wkhtmltopdf，回退 LaTeX
            extra_args.extend(["--pdf-engine=xelatex", "-V", "CJKmainfont=SimSun"])
        elif output_format == "docx":
            pass
        elif output_format == "html":
            extra_args.append("--standalone")

        pypandoc.convert_file(
            input_path,
            output_format,
            outputfile=output_path,
            extra_args=extra_args if extra_args else None,
        )

        return [output_path]

    # ---- 图片 → PDF -------------------------------------------------------

    def _image_to_pdf(
        self,
        img_path: str,
        progress_cb: Optional[Callable[[int, str], None]],
    ) -> list[str]:
        if "PyMuPDF" in _import_errors:
            raise ConversionError("PyMuPDF 库未安装，无法执行 图片→PDF 转换")

        pdf_path = self._get_output_path(img_path, ".pdf")
        self._emit(progress_cb, 30, "正在将图片转为 PDF…")

        # 用 PyMuPDF 打开图片并转为一页 PDF
        img_doc = fitz.open(img_path)
        # convert_to_pdf 将图片转为 PDF 字节流
        pdf_bytes = img_doc.convert_to_pdf()
        img_doc.close()

        out_doc = fitz.open("pdf", pdf_bytes)
        out_doc.save(pdf_path, deflate=True, garbage=4)
        out_doc.close()

        return [pdf_path]

    # ---- 辅助 -------------------------------------------------------------

    @staticmethod
    def _emit(
        cb: Optional[Callable[[int, str], None]],
        percent: int,
        message: str,
    ) -> None:
        if cb:
            cb(percent, message)

    @staticmethod
    def get_missing_deps() -> dict[str, str]:
        """返回缺失的依赖及其错误信息。"""
        return dict(_import_errors)
